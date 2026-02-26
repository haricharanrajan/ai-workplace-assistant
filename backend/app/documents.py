from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from sqlalchemy.orm import Session
import shutil
import os

from pypdf import PdfReader
from docx import Document as DocxDocument

from app.deps import get_db, require_admin
from app import models
from app.rag import embed_document


router = APIRouter(
    prefix="/documents",
    tags=["Documents"]
)

# Folder to store uploaded files
UPLOAD_FOLDER = "uploaded_docs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Text Extraction Helpers

def extract_text_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()

def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)

    return "\n".join(parts).strip()

def extract_text_from_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(parts).strip()

# Upload Endpoint (Admin Only)

@router.post("/upload")
def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_admin=Depends(require_admin)
):
    # Upload and index a document.
    # Only admins are authorized.

    # File type validation

    allowed_extensions = [".txt", ".pdf", ".docx"]
    ext = os.path.splitext(file.filename)[1].lower()

    if ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Allowed: txt, pdf, docx"
        )

    # Save file to disk

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Extract text

    try:
        if ext == ".txt":
            text = extract_text_from_txt(file_path)

        elif ext == ".pdf":
            text = extract_text_from_pdf(file_path)

        elif ext == ".docx":
            text = extract_text_from_docx(file_path)

        else:
            text = ""

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Text extraction failed: {str(e)}"
        )

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="No readable text found in the document."
        )

    # Generate embeddings

    try:
        embed_document(text)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Embedding generation failed: {str(e)}"
        )

    # Save metadata to DB

    new_doc = models.Document(
        filename=file.filename,
        filepath=file_path
    )

    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    return {
        "message": "Document uploaded and indexed successfully",
        "filename": file.filename
    }